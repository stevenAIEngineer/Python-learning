from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture(
    'me.png', 
    width=Inches(1.0)
)

# name, phone number, email details
name = input('What is your name? ')
speak('Hello ' + name + ', welcome to the CV generator')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

# about me
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself? ')
)

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company: ')
start_date = input('From Date: ')
end_date = input('To Date: ')

p.add_run(company + ' ').bold = True
p.add_run(start_date + ' - ' + end_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company  + ' ')
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? (yes/no) ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company: ')
        start_date = input('From Date: ')
        end_date = input('To Date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(start_date + ' - ' + end_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company  + ': ')
        p.add_run(experience_details)
    else:
        break

# skills
document.add_heading('Skills')
skill = input('Enter skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'
while True:
     has_more_skills = input(
            'Do you have more skills? (yes/no) ')
     if has_more_skills.lower() == 'yes':
            skill = input('Enter skill: ')
            p = document.add_paragraph(skill)
            p.style = 'List Bullet'
     else:
            break
        
# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated in Python using Visual Studio Code"

document.save('cv.docx')
